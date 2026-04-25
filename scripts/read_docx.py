
from docx import Document

doc = Document('党旗飘飘题库.docx')

print("=== 文档内容 ===\n")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"[{i}] {para.text}")

print("\n=== 文档表格 ===\n")
for table_idx, table in enumerate(doc.tables):
    print(f"表格 {table_idx}:")
    for row_idx, row in enumerate(table.rows):
        row_data = [cell.text.strip() for cell in row.cells]
        print(f"  行 {row_idx}: {row_data}")
