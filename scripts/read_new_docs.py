
from docx import Document
import os

doc_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', '第四套-第六套.docx')
doc = Document(doc_path)

print("=== 文档前50段内容 ===\n")
count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"[{i}] {repr(text)}")  # 用repr看特殊字符
        count += 1
        if count >= 50:
            break

print("\n\n=== 文档表格内容 ===\n")
for table_idx, table in enumerate(doc.tables):
    print(f"\n--- 表格 {table_idx} ---")
    for row_idx, row in enumerate(table.rows):
        row_data = [cell.text.strip() for cell in row.cells]
        print(f"行 {row_idx}: {row_data}")
